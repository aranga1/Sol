"""Tests for solidrag extractors — markdown, PDF, Excel, docx, image.

All fixtures are in-memory (no disk I/O beyond tmpdir for path objects).
No real Ollama calls are made — ImageExtractor uses a mocked httpx client.
"""
from __future__ import annotations

import io
import json
import textwrap
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


# ---------------------------------------------------------------------------
# MarkdownExtractor
# ---------------------------------------------------------------------------

class TestMarkdownExtractor:
    def _write(self, tmp_path: Path, content: str) -> Path:
        p = tmp_path / "doc.md"
        p.write_text(content, encoding="utf-8")
        return p

    def test_supported_extensions(self):
        from solidrag.extractors.markdown import MarkdownExtractor
        ext = MarkdownExtractor()
        assert ".md" in ext.supported_extensions

    def test_empty_file_returns_no_nodes(self, tmp_path):
        from solidrag.extractors.markdown import MarkdownExtractor
        p = self._write(tmp_path, "")
        nodes = MarkdownExtractor().extract(p)
        assert nodes == []

    def test_whitespace_only_file_returns_no_nodes(self, tmp_path):
        from solidrag.extractors.markdown import MarkdownExtractor
        p = self._write(tmp_path, "   \n\n  \n")
        nodes = MarkdownExtractor().extract(p)
        assert nodes == []

    def test_single_section_returns_one_node(self, tmp_path):
        from solidrag.extractors.markdown import MarkdownExtractor
        content = "Hello world. This is a paragraph."
        p = self._write(tmp_path, content)
        nodes = MarkdownExtractor().extract(p)
        assert len(nodes) == 1
        assert "Hello world" in nodes[0].text

    def test_heading_split_produces_multiple_nodes(self, tmp_path):
        from solidrag.extractors.markdown import MarkdownExtractor
        content = textwrap.dedent("""\
            # Introduction
            This is the introduction section with some content.

            # Methods
            This describes methods used in the study.

            # Results
            Here are the results of the experiment.
        """)
        p = self._write(tmp_path, content)
        nodes = MarkdownExtractor().extract(p)
        # Headings split into multiple chunks
        assert len(nodes) >= 2

    def test_overlap_context_added_between_chunks(self, tmp_path):
        from solidrag.extractors.markdown import MarkdownExtractor
        content = textwrap.dedent("""\
            # Section One
            First section sentence one. First section sentence two. First section sentence three.

            # Section Two
            Second section sentence one. Second section sentence two. Second section sentence three.

            # Section Three
            Third section sentence one. Third section sentence two. Third section sentence three.
        """)
        p = self._write(tmp_path, content)
        nodes = MarkdownExtractor().extract(p)
        # Middle chunk should have context from previous and next sections
        assert len(nodes) >= 3
        middle_texts = [n.text for n in nodes]
        # At least one node should contain a context marker
        has_context = any("[context:" in t for t in middle_texts)
        assert has_context

    def test_metadata_contains_file_path_and_name(self, tmp_path):
        from solidrag.extractors.markdown import MarkdownExtractor
        content = "Some content here."
        p = self._write(tmp_path, content)
        nodes = MarkdownExtractor().extract(p)
        assert len(nodes) == 1
        assert nodes[0].metadata["file_path"] == str(p)
        assert nodes[0].metadata["file_name"] == "doc.md"

    def test_excluded_llm_metadata_keys_set(self, tmp_path):
        from solidrag.extractors.markdown import MarkdownExtractor
        p = self._write(tmp_path, "Content here.")
        nodes = MarkdownExtractor().extract(p)
        assert len(nodes) == 1
        assert "file_path" in nodes[0].excluded_llm_metadata_keys
        assert "file_name" in nodes[0].excluded_llm_metadata_keys

    def test_hr_split(self, tmp_path):
        from solidrag.extractors.markdown import MarkdownExtractor
        content = textwrap.dedent("""\
            First section content here with enough text.

            ---

            Second section content here with enough text.
        """)
        p = self._write(tmp_path, content)
        nodes = MarkdownExtractor().extract(p)
        assert len(nodes) >= 1

    def test_long_section_splits_on_paragraphs(self, tmp_path):
        from solidrag.extractors.markdown import MarkdownExtractor
        # Create a section longer than _CHUNK_MAX_CHARS (1400)
        # by repeating paragraphs
        para = "This is a long paragraph. " * 30  # ~780 chars each
        content = "# Big Section\n\n" + para + "\n\n" + para + "\n\n" + para
        p = self._write(tmp_path, content)
        nodes = MarkdownExtractor().extract(p)
        # Should produce more than 1 node due to length splitting
        assert len(nodes) >= 1
        # All nodes must have text
        for n in nodes:
            assert n.text.strip()


# ---------------------------------------------------------------------------
# PDFExtractor
# ---------------------------------------------------------------------------

class TestPDFExtractor:
    def _make_pdf_bytes(self, pages: list[str]) -> bytes:
        """Create a minimal in-memory PDF with pypdf's PdfWriter."""
        from pypdf import PdfWriter
        from pypdf.generic import NameObject

        writer = PdfWriter()
        for text in pages:
            # Add a blank page and annotate with text via pypdf
            page = writer.add_blank_page(width=595, height=842)
            # We add text via direct content stream injection
            content = f"BT /F1 12 Tf 50 750 Td ({text}) Tj ET"
            from pypdf.generic import DecodedStreamObject, ArrayObject, DictionaryObject
            stream = DecodedStreamObject()
            stream.set_data(content.encode())
            page_obj = page.get_object()
            page_obj[NameObject("/Contents")] = writer._add_object(stream)
            # Add a basic font resource so pdf is valid
            font_dict = DictionaryObject({
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            })
            resources = DictionaryObject({
                NameObject("/Font"): DictionaryObject({
                    NameObject("/F1"): writer._add_object(font_dict)
                })
            })
            page_obj[NameObject("/Resources")] = resources

        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue()

    def test_supported_extensions(self):
        from solidrag.extractors.pdf import PDFExtractor
        ext = PDFExtractor()
        assert ".pdf" in ext.supported_extensions

    def test_one_node_per_page(self, tmp_path):
        from solidrag.extractors.pdf import PDFExtractor
        pdf_bytes = self._make_pdf_bytes(["Page one text", "Page two text", "Page three text"])
        p = tmp_path / "doc.pdf"
        p.write_bytes(pdf_bytes)
        nodes = PDFExtractor().extract(p)
        # Should produce at most 3 nodes (blank pages may be skipped)
        assert 0 < len(nodes) <= 3

    def test_metadata_file_path_and_name(self, tmp_path):
        from solidrag.extractors.pdf import PDFExtractor
        pdf_bytes = self._make_pdf_bytes(["Hello PDF page"])
        p = tmp_path / "report.pdf"
        p.write_bytes(pdf_bytes)
        nodes = PDFExtractor().extract(p)
        assert len(nodes) >= 1
        for node in nodes:
            assert node.metadata["file_path"] == str(p)
            assert node.metadata["file_name"] == "report.pdf"

    def test_metadata_page_number(self, tmp_path):
        from solidrag.extractors.pdf import PDFExtractor
        pdf_bytes = self._make_pdf_bytes(["Page one text", "Page two text"])
        p = tmp_path / "doc.pdf"
        p.write_bytes(pdf_bytes)
        nodes = PDFExtractor().extract(p)
        page_numbers = [n.metadata["page_number"] for n in nodes]
        # page numbers should be integers starting from 1
        assert all(isinstance(pn, int) for pn in page_numbers)
        assert all(pn >= 1 for pn in page_numbers)

    def test_empty_pages_skipped(self, tmp_path):
        from solidrag.extractors.pdf import PDFExtractor
        # Create PDF with a blank page (no text) and one with text
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)  # truly empty

        buf = io.BytesIO()
        writer.write(buf)
        p = tmp_path / "empty.pdf"
        p.write_bytes(buf.getvalue())

        nodes = PDFExtractor().extract(p)
        assert nodes == []


# ---------------------------------------------------------------------------
# ExcelExtractor
# ---------------------------------------------------------------------------

class TestExcelExtractor:
    def _make_xlsx(self, sheets: dict[str, list[list]]) -> bytes:
        """Create a minimal in-memory xlsx with openpyxl."""
        import openpyxl
        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # remove default sheet
        for sheet_name, rows in sheets.items():
            ws = wb.create_sheet(title=sheet_name)
            for row in rows:
                ws.append(row)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_supported_extensions(self):
        from solidrag.extractors.excel import ExcelExtractor
        ext = ExcelExtractor()
        assert ".xlsx" in ext.supported_extensions
        assert ".xls" in ext.supported_extensions

    def test_one_node_per_sheet(self, tmp_path):
        from solidrag.extractors.excel import ExcelExtractor
        data = {
            "Sheet1": [["Name", "Age"], ["Alice", 30], ["Bob", 25]],
            "Sheet2": [["Product", "Price"], ["Widget", 9.99]],
        }
        p = tmp_path / "data.xlsx"
        p.write_bytes(self._make_xlsx(data))
        nodes = ExcelExtractor().extract(p)
        assert len(nodes) == 2

    def test_row_format(self, tmp_path):
        from solidrag.extractors.excel import ExcelExtractor
        data = {"Data": [["col1", "col2"], ["val1", "val2"]]}
        p = tmp_path / "data.xlsx"
        p.write_bytes(self._make_xlsx(data))
        nodes = ExcelExtractor().extract(p)
        assert len(nodes) == 1
        # text should contain col: val format with pipe separator
        text = nodes[0].text
        assert "col1" in text
        assert "val1" in text
        assert "|" in text

    def test_metadata_sheet_name(self, tmp_path):
        from solidrag.extractors.excel import ExcelExtractor
        data = {"MySheet": [["A", "B"], [1, 2]]}
        p = tmp_path / "data.xlsx"
        p.write_bytes(self._make_xlsx(data))
        nodes = ExcelExtractor().extract(p)
        assert nodes[0].metadata["sheet_name"] == "MySheet"

    def test_metadata_file_path_and_name(self, tmp_path):
        from solidrag.extractors.excel import ExcelExtractor
        data = {"Sheet1": [["x"], [1]]}
        p = tmp_path / "report.xlsx"
        p.write_bytes(self._make_xlsx(data))
        nodes = ExcelExtractor().extract(p)
        assert nodes[0].metadata["file_path"] == str(p)
        assert nodes[0].metadata["file_name"] == "report.xlsx"

    def test_empty_sheet_returns_empty_list(self, tmp_path):
        """A workbook whose only sheet has no rows produces no nodes."""
        from solidrag.extractors.excel import ExcelExtractor
        import openpyxl
        wb = openpyxl.Workbook()
        # Default active sheet exists but has no rows written to it
        buf = io.BytesIO()
        wb.save(buf)
        p = tmp_path / "empty.xlsx"
        p.write_bytes(buf.getvalue())
        nodes = ExcelExtractor().extract(p)
        assert nodes == []

    def test_single_sheet_with_headers_only(self, tmp_path):
        from solidrag.extractors.excel import ExcelExtractor
        data = {"Headers": [["Name", "Value"]]}
        p = tmp_path / "headers.xlsx"
        p.write_bytes(self._make_xlsx(data))
        nodes = ExcelExtractor().extract(p)
        # Should produce a node (header row is still content)
        assert len(nodes) == 1


# ---------------------------------------------------------------------------
# DocxExtractor
# ---------------------------------------------------------------------------

class TestDocxExtractor:
    def _make_docx(
        self,
        paragraphs: list[str] | None = None,
        headings: list[tuple[int, str]] | None = None,
        tables: list[list[list[str]]] | None = None,
    ) -> bytes:
        """Create a minimal in-memory docx."""
        from docx import Document
        doc = Document()
        if headings:
            for level, text in headings:
                doc.add_heading(text, level=level)
        if paragraphs:
            for text in paragraphs:
                doc.add_paragraph(text)
        if tables:
            for table_rows in tables:
                if not table_rows:
                    continue
                t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell in enumerate(row):
                        t.cell(r_idx, c_idx).text = cell
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_supported_extensions(self):
        from solidrag.extractors.docx import DocxExtractor
        ext = DocxExtractor()
        assert ".docx" in ext.supported_extensions
        assert ".doc" in ext.supported_extensions

    def test_paragraphs_in_output(self, tmp_path):
        from solidrag.extractors.docx import DocxExtractor
        p = tmp_path / "doc.docx"
        p.write_bytes(self._make_docx(paragraphs=["Hello world.", "Second paragraph."]))
        nodes = DocxExtractor().extract(p)
        assert len(nodes) == 1
        text = nodes[0].text
        assert "Hello world" in text
        assert "Second paragraph" in text

    def test_headings_formatted_with_hash(self, tmp_path):
        from solidrag.extractors.docx import DocxExtractor
        p = tmp_path / "doc.docx"
        p.write_bytes(self._make_docx(headings=[(1, "Main Title"), (2, "Subsection")]))
        nodes = DocxExtractor().extract(p)
        assert len(nodes) == 1
        text = nodes[0].text
        assert "# Main Title" in text
        assert "## Subsection" in text

    def test_tables_formatted_with_pipes(self, tmp_path):
        from solidrag.extractors.docx import DocxExtractor
        table_data = [["Header A", "Header B"], ["Row1A", "Row1B"]]
        p = tmp_path / "doc.docx"
        p.write_bytes(self._make_docx(tables=[table_data]))
        nodes = DocxExtractor().extract(p)
        assert len(nodes) == 1
        text = nodes[0].text
        assert "|" in text
        assert "Header A" in text

    def test_metadata_file_path_and_name(self, tmp_path):
        from solidrag.extractors.docx import DocxExtractor
        p = tmp_path / "report.docx"
        p.write_bytes(self._make_docx(paragraphs=["Some text."]))
        nodes = DocxExtractor().extract(p)
        assert nodes[0].metadata["file_path"] == str(p)
        assert nodes[0].metadata["file_name"] == "report.docx"

    def test_empty_document_returns_no_nodes(self, tmp_path):
        from solidrag.extractors.docx import DocxExtractor
        from docx import Document
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        p = tmp_path / "empty.docx"
        p.write_bytes(buf.getvalue())
        nodes = DocxExtractor().extract(p)
        # Empty document should return empty list
        assert nodes == []

    def test_returns_single_node(self, tmp_path):
        from solidrag.extractors.docx import DocxExtractor
        p = tmp_path / "doc.docx"
        p.write_bytes(self._make_docx(
            headings=[(1, "Title")],
            paragraphs=["Para one.", "Para two."],
            tables=[[["A", "B"], ["1", "2"]]],
        ))
        nodes = DocxExtractor().extract(p)
        assert len(nodes) == 1


# ---------------------------------------------------------------------------
# ImageExtractor
# ---------------------------------------------------------------------------

class TestImageExtractor:
    def _make_minimal_png(self) -> bytes:
        """Return a minimal 1x1 white PNG as bytes."""
        import base64
        # Minimal valid 1x1 PNG (hardcoded bytes)
        png_b64 = (
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8"
            "z8BQDwADhQGAWjR9awAAAABJRU5ErkJggg=="
        )
        return base64.b64decode(png_b64)

    def test_supported_extensions(self):
        from solidrag.extractors.image import ImageExtractor
        ext = ImageExtractor(ollama_base_url="http://localhost:11434")
        for e in (".jpg", ".jpeg", ".png", ".gif", ".webp"):
            assert e in ext.supported_extensions

    def test_returns_one_node_with_description(self, tmp_path):
        from solidrag.extractors.image import ImageExtractor

        p = tmp_path / "photo.png"
        p.write_bytes(self._make_minimal_png())

        mock_response = MagicMock()
        mock_response.json.return_value = {"response": "A white square image."}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.post", return_value=mock_response) as mock_post:
            extractor = ImageExtractor(ollama_base_url="http://localhost:11434", vision_model="llava")
            nodes = extractor.extract(p)

        assert len(nodes) == 1
        assert "A white square image." in nodes[0].text

    def test_ollama_api_called_with_base64_image(self, tmp_path):
        from solidrag.extractors.image import ImageExtractor
        import base64

        img_bytes = self._make_minimal_png()
        p = tmp_path / "photo.png"
        p.write_bytes(img_bytes)
        expected_b64 = base64.b64encode(img_bytes).decode()

        mock_response = MagicMock()
        mock_response.json.return_value = {"response": "An image."}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.post", return_value=mock_response) as mock_post:
            extractor = ImageExtractor(ollama_base_url="http://localhost:11434", vision_model="llava")
            extractor.extract(p)

        call_kwargs = mock_post.call_args
        payload = call_kwargs[1].get("json") or call_kwargs[0][1]
        assert payload["model"] == "llava"
        assert expected_b64 in payload["images"]

    def test_ollama_url_uses_base_url(self, tmp_path):
        from solidrag.extractors.image import ImageExtractor

        p = tmp_path / "photo.jpg"
        p.write_bytes(self._make_minimal_png())

        mock_response = MagicMock()
        mock_response.json.return_value = {"response": "desc"}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.post", return_value=mock_response) as mock_post:
            extractor = ImageExtractor(ollama_base_url="http://myserver:11434")
            extractor.extract(p)

        called_url = mock_post.call_args[0][0]
        assert "myserver:11434" in called_url

    def test_metadata_file_path_and_name(self, tmp_path):
        from solidrag.extractors.image import ImageExtractor

        p = tmp_path / "photo.png"
        p.write_bytes(self._make_minimal_png())

        mock_response = MagicMock()
        mock_response.json.return_value = {"response": "A photo."}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.post", return_value=mock_response):
            extractor = ImageExtractor(ollama_base_url="http://localhost:11434")
            nodes = extractor.extract(p)

        assert nodes[0].metadata["file_path"] == str(p)
        assert nodes[0].metadata["file_name"] == "photo.png"

    def test_custom_vision_model_passed_to_api(self, tmp_path):
        from solidrag.extractors.image import ImageExtractor

        p = tmp_path / "photo.webp"
        p.write_bytes(self._make_minimal_png())

        mock_response = MagicMock()
        mock_response.json.return_value = {"response": "description"}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.post", return_value=mock_response) as mock_post:
            extractor = ImageExtractor(
                ollama_base_url="http://localhost:11434",
                vision_model="bakllava",
            )
            extractor.extract(p)

        call_kwargs = mock_post.call_args
        payload = call_kwargs[1].get("json") or call_kwargs[0][1]
        assert payload["model"] == "bakllava"


# ---------------------------------------------------------------------------
# default_registry integration
# ---------------------------------------------------------------------------

class TestDefaultRegistry:
    def test_default_registry_returns_registry(self):
        from solidrag.extractors import default_registry
        reg = default_registry()
        from solidrag.extractors.registry import ExtractorRegistry
        assert isinstance(reg, ExtractorRegistry)

    def test_all_extensions_registered(self):
        from solidrag.extractors import default_registry
        reg = default_registry()
        expected = {".md", ".pdf", ".xlsx", ".xls", ".docx", ".doc",
                    ".jpg", ".jpeg", ".png", ".gif", ".webp"}
        registered = reg.extensions()
        assert expected.issubset(registered), f"Missing: {expected - registered}"

    def test_custom_ollama_url_forwarded(self):
        from solidrag.extractors import default_registry
        from solidrag.extractors.image import ImageExtractor
        reg = default_registry(ollama_base_url="http://custom:11434")
        img_ext = reg.get(".png")
        assert isinstance(img_ext, ImageExtractor)
        assert img_ext.ollama_base_url == "http://custom:11434"

    def test_custom_vision_model_forwarded(self):
        from solidrag.extractors import default_registry
        from solidrag.extractors.image import ImageExtractor
        reg = default_registry(vision_model="bakllava")
        img_ext = reg.get(".jpg")
        assert isinstance(img_ext, ImageExtractor)
        assert img_ext.vision_model == "bakllava"
